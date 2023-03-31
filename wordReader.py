import json
import os
from docx import Document

def extract_text(doc):
    start_idx, end_idx = -1, -1
    
    for i, para in enumerate(doc.paragraphs):
        if 'RESUMEN EJECUTIVO' in para.text:
            start_idx = i
        elif start_idx >= 0 and 'A continuación' in para.text:
            end_idx = i
            break

    if start_idx >= 0 and end_idx >= 0:
        text = '\n'.join([para.text for para in doc.paragraphs[start_idx+1:end_idx]])
        return text.strip()
    else:
        return None

file = "C:\\Users\\lorni\\Documents\\Desarrollo\\prueba_lectura_archivos_doc\\informe_unix_con_comentarios2.docx"

# Obtiene el nombre base del archivo de Word
file_name = os.path.splitext(os.path.basename(file))[0]

# Lee el documento de Word
doc = Document(file)

# Extrae el texto del resumen ejecutivo
resumen_ejecutivo = extract_text(doc)

# Busca todas las tablas en el documento y convierte su contenido en una lista de listas
tables = []
for table in doc.tables:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        table_data.append(row_data)
    tables.append(table_data)

# Combina el texto del resumen ejecutivo y las tablas en un objeto JSON y lo guarda en un archivo con el mismo nombre que el archivo de Word
data = {"resumen ejecutivo": resumen_ejecutivo, "tablas": tables}
with open(file_name + '.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False)
